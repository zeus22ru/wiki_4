#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Структурное разбиение документов на чанки (заголовки, списки, абзацы, таблицы).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from bs4 import BeautifulSoup, NavigableString, Tag

from config import settings, get_logger

logger = get_logger(__name__)


def chunk_text_fixed_size(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
    """Классическое разбиение по размеру с overlap (как раньше в create_vector_db)."""
    if chunk_size is None:
        chunk_size = settings.CHUNK_SIZE
    if overlap is None:
        overlap = settings.CHUNK_OVERLAP
    chunks: List[str] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = start + chunk_size
        chunk = text[start:end]

        if end < text_len:
            last_period = chunk.rfind(".")
            last_question = chunk.rfind("?")
            last_exclamation = chunk.rfind("!")
            last_boundary = max(last_period, last_question, last_exclamation)

            if last_boundary > chunk_size // 2:
                chunk = text[start : start + last_boundary + 1]
                end = start + last_boundary + 1

        chunks.append(chunk.strip())
        start = end - overlap

    return [c for c in chunks if len(c) > 50]


def _join_path(headings: List[str]) -> str:
    cleaned = [h.strip() for h in headings if h and str(h).strip()]
    return " → ".join(cleaned) if cleaned else ""


def _merge_small_blocks(
    blocks: List[Dict[str, Any]],
    max_chars: int,
    min_chars: int,
) -> List[Dict[str, Any]]:
    """Склеить слишком короткие блоки с соседями в пределах max_chars."""
    if not blocks:
        return []
    merged: List[Dict[str, Any]] = []
    buf_text: List[str] = []
    buf_meta: Optional[Dict[str, Any]] = None

    def flush() -> None:
        nonlocal buf_text, buf_meta
        if not buf_text or buf_meta is None:
            buf_text = []
            buf_meta = None
            return
        text = "\n\n".join(t for t in buf_text if t).strip()
        if len(text) >= min_chars or not merged:
            merged.append({
                "text": text,
                "section_path": buf_meta.get("section_path", ""),
                "chunk_kind": buf_meta.get("chunk_kind", "text"),
                "parent_headings": buf_meta.get("parent_headings", []),
            })
        else:
            prev = merged[-1]
            prev["text"] = (prev["text"] + "\n\n" + text).strip()
        buf_text = []
        buf_meta = None

    for block in blocks:
        t = (block.get("text") or "").strip()
        if not t:
            continue
        if buf_meta is None:
            buf_meta = {
                "section_path": block.get("section_path", ""),
                "chunk_kind": block.get("chunk_kind", "text"),
                "parent_headings": list(block.get("parent_headings") or []),
            }
            buf_text = [t]
            continue
        candidate = "\n\n".join(buf_text + [t])
        same_section = buf_meta.get("section_path") == block.get("section_path")
        if same_section and len(candidate) <= max_chars:
            buf_text.append(t)
        else:
            flush()
            buf_meta = {
                "section_path": block.get("section_path", ""),
                "chunk_kind": block.get("chunk_kind", "text"),
                "parent_headings": list(block.get("parent_headings") or []),
            }
            buf_text = [t]
    flush()
    return [m for m in merged if len(m.get("text") or "") >= min_chars]


def chunk_html_structural(html_content: str, page_title: str, relative_path: str) -> List[Dict[str, Any]]:
    """Разбиение HTML по заголовкам и блочным элементам."""
    soup = BeautifulSoup(html_content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    page_heading = page_title or ""
    outline: List[str] = []

    def current_section_path() -> str:
        parts = [page_heading] if page_heading else []
        parts.extend(h for h in outline if h)
        return _join_path(parts)

    def current_parent_headings() -> List[str]:
        parts = [page_heading] if page_heading else []
        parts.extend(h for h in outline if h)
        return parts

    blocks: List[Dict[str, Any]] = []

    def heading_level(tag_name: str) -> Optional[int]:
        if tag_name and tag_name.startswith("h") and len(tag_name) == 2 and tag_name[1].isdigit():
            return int(tag_name[1])
        return None

    def walk(parent: Tag) -> None:
        nonlocal outline
        for child in parent.children:
            if isinstance(child, NavigableString):
                txt = str(child).strip()
                if txt:
                    ph = current_parent_headings()
                    blocks.append({
                        "text": txt,
                        "section_path": current_section_path(),
                        "chunk_kind": "text",
                        "parent_headings": ph,
                    })
                continue
            if not isinstance(child, Tag):
                continue
            name = child.name.lower()
            lvl = heading_level(name)
            if lvl is not None:
                title = child.get_text(separator=" ", strip=True)
                outline = outline[: lvl - 1]
                outline.append(title or f"Раздел {lvl}")
                if title:
                    ph = current_parent_headings()
                    blocks.append({
                        "text": title,
                        "section_path": current_section_path(),
                        "chunk_kind": "heading",
                        "parent_headings": ph,
                    })
                walk(child)
                continue
            if name in ("ul", "ol"):
                items = []
                for li in child.find_all("li", recursive=False):
                    lit = li.get_text(separator=" ", strip=True)
                    if lit:
                        items.append(f"• {lit}" if name == "ul" else lit)
                if items:
                    ph = current_parent_headings()
                    blocks.append({
                        "text": "\n".join(items),
                        "section_path": current_section_path(),
                        "chunk_kind": "list",
                        "parent_headings": ph,
                    })
                continue
            if name == "table":
                rows_out: List[str] = []
                for tr in child.find_all("tr"):
                    cells = tr.find_all(["th", "td"])
                    row = " | ".join(c.get_text(separator=" ", strip=True) for c in cells)
                    if row.strip():
                        rows_out.append(row)
                if rows_out:
                    ph = current_parent_headings()
                    blocks.append({
                        "text": "\n".join(rows_out),
                        "section_path": current_section_path(),
                        "chunk_kind": "table",
                        "parent_headings": ph,
                    })
                continue
            if name in ("p", "div", "section", "article", "blockquote", "pre"):
                txt = child.get_text(separator=" ", strip=True)
                if txt:
                    kind = "code" if name == "pre" else "text"
                    ph = current_parent_headings()
                    blocks.append({
                        "text": txt,
                        "section_path": current_section_path(),
                        "chunk_kind": kind,
                        "parent_headings": ph,
                    })
                else:
                    walk(child)
                continue
            walk(child)

    body = soup.body or soup
    walk(body)

    return _merge_small_blocks(
        blocks,
        settings.STRUCTURAL_CHUNK_MAX_CHARS,
        settings.STRUCTURAL_CHUNK_MIN_CHARS,
    )


def chunk_plain_paragraphs(text: str, title: str, relative_path: str) -> List[Dict[str, Any]]:
    """Запасной путь: абзацы и псевдо-заголовки."""
    text = re.sub(r"[ \t]+", " ", text or "").strip()
    if not text:
        return []
    parts = re.split(r"\n\s*\n+", text)
    blocks: List[Dict[str, Any]] = []
    current_heading = title or Path(relative_path).stem

    for part in parts:
        p = part.strip()
        if not p:
            continue
        lines = p.split("\n")
        first = lines[0].strip()
        is_heading = (
            (len(first) < 120 and len(lines) > 1 and first.isupper())
            or (len(first) < 100 and first.endswith(":") and len(lines) > 1)
        )
        if is_heading:
            current_heading = first.rstrip(":")
            rest = "\n".join(lines[1:]).strip()
            if rest:
                path = _join_path([title, current_heading] if title else [current_heading])
                blocks.append({
                    "text": rest,
                    "section_path": path,
                    "chunk_kind": "text",
                    "parent_headings": [title, current_heading] if title else [current_heading],
                })
            continue
        path = _join_path([title, current_heading] if title else [current_heading])
        blocks.append({
            "text": p,
            "section_path": path,
            "chunk_kind": "text",
            "parent_headings": [title, current_heading] if title else [current_heading],
        })

    return _merge_small_blocks(
        blocks,
        settings.STRUCTURAL_CHUNK_MAX_CHARS,
        settings.STRUCTURAL_CHUNK_MIN_CHARS,
    )


def chunk_docx_structural(doc_path: Path, title: str, relative_path: str) -> Optional[List[Dict[str, Any]]]:
    try:
        from docx import Document
    except ImportError:
        return None
    try:
        doc = Document(doc_path)
    except Exception as e:
        logger.error("DOCX structural: %s", e)
        return None

    blocks: List[Dict[str, Any]] = []
    headings: List[str] = [title] if title else []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        style = (para.style.name if para.style else "") or ""
        if style.startswith("Heading"):
            digits = "".join(ch for ch in style if ch.isdigit())
            level = int(digits) if digits else 1
            while len(headings) < level:
                headings.append("")
            headings = headings[:level]
            headings[level - 1] = t
            blocks.append({
                "text": t,
                "section_path": _join_path(headings),
                "chunk_kind": "heading",
                "parent_headings": list(headings),
            })
        else:
            blocks.append({
                "text": t,
                "section_path": _join_path(headings),
                "chunk_kind": "text",
                "parent_headings": list(headings),
            })

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                row_text = " | ".join(cells)
                blocks.append({
                    "text": row_text,
                    "section_path": _join_path(headings),
                    "chunk_kind": "table_row",
                    "parent_headings": list(headings),
                })

    return _merge_small_blocks(
        blocks,
        settings.STRUCTURAL_CHUNK_MAX_CHARS,
        settings.STRUCTURAL_CHUNK_MIN_CHARS,
    )


def build_chunks_for_file(
    file_path: Path,
    doc_data: Dict[str, str],
    extract_func: Callable[[Path], Optional[Dict[str, str]]],
) -> List[Dict[str, Any]]:
    """
    Построить список чанков с полями text, section_path, chunk_kind, parent_headings (JSON-serializable).
    """
    ext = file_path.suffix.lower()
    rel = doc_data.get("path") or file_path.name
    title = doc_data.get("title") or file_path.stem
    content = doc_data.get("content") or ""

    if not settings.STRUCTURAL_CHUNKING_ENABLED:
        raw_chunks = chunk_text_fixed_size(content)
        return [
            {
                "text": c,
                "section_path": title,
                "chunk_kind": "text",
                "parent_headings": [title],
            }
            for c in raw_chunks
        ]

    chunks: List[Dict[str, Any]] = []
    if ext in (".html", ".htm"):
        try:
            html_src = file_path.read_text(encoding="utf-8", errors="replace")
            chunks = chunk_html_structural(html_src, title, rel)
        except Exception as e:
            logger.warning("HTML structural fallback: %s", e)
        if not chunks:
            chunks = chunk_plain_paragraphs(content, title, rel)
    elif ext == ".docx":
        structured = chunk_docx_structural(file_path, title, rel)
        chunks = structured or chunk_plain_paragraphs(content, title, rel)
    else:
        chunks = chunk_plain_paragraphs(content, title, rel)

    return chunks
